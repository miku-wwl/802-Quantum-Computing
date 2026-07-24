"""Generate the comprehensive Task 4 Word report and its fixed layout."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_task3_report import (
    BLUE,
    DARK_BLUE,
    MUTED,
    NAVY,
    add_body,
    add_bullet,
    add_callout,
    add_caption,
    add_linked_section,
    add_page_number,
    add_section_heading,
    add_table,
    set_run_font,
    set_section_geometry,
)


ROOT = Path(__file__).resolve().parents[1]
TASK_DIR = ROOT / "submission" / "Task_4_Quantum_ML"
OUTPUT = TASK_DIR / "Task_4_Quantum_ML_Report.docx"


def configure_document(document: Document) -> None:
    """Apply the standard_business_brief preset and Task 4 page furniture."""

    document.settings.odd_and_even_pages_header_footer = False
    section = document.sections[0]
    set_section_geometry(section)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style_name in ("List Bullet", "List Number"):
        style = styles[list_style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = MUTED
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(5)
    caption.paragraph_format.keep_with_next = True

    for header_part in (
        section.header,
        section.even_page_header,
        section.first_page_header,
    ):
        header = header_part.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_run = header.add_run("MSE802  •  ASSESSMENT 2  •  TASK 4")
        set_run_font(header_run, size=8.5, color=MUTED, bold=True)

    for footer_part in (
        section.footer,
        section.even_page_footer,
        section.first_page_footer,
    ):
        footer = footer_part.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("Quantum ML Analysis  |  Page ")
        set_run_font(footer_run, size=8.5, color=MUTED)
        add_page_number(footer)

    document.core_properties.title = "Quantum Machine Learning — Task 4 Report"
    document.core_properties.subject = "MSE802 Assessment 2"
    document.core_properties.author = "MSE802 student submission"
    document.core_properties.keywords = (
        "quantum machine learning, Qiskit, Aer, Quokka, SPSA, logistic regression"
    )


def add_cover(document: Document) -> None:
    """Use the documents-skill editorial_cover pattern."""

    for _ in range(5):
        document.add_paragraph()
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("TECHNICAL ANALYSIS REPORT")
    set_run_font(run, size=10, color=BLUE, bold=True)
    kicker.paragraph_format.space_after = Pt(18)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Quantum Machine Learning")
    set_run_font(run, size=30, color=NAVY, bold=True)
    title.paragraph_format.space_after = Pt(8)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Circuit and input analysis, optimization evidence,\n"
        "and a fully classical no-circuit comparison"
    )
    set_run_font(run, size=15, color=DARK_BLUE)
    subtitle.paragraph_format.space_after = Pt(30)

    context = document.add_paragraph()
    context.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = context.add_run("MSE802 Quantum Computing  •  Assessment 2  •  Task 4")
    set_run_font(run, size=11, color=MUTED, bold=True)
    context.paragraph_format.space_after = Pt(98)

    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = metadata.add_run(
        "Local implementation: Python 3.11, Qiskit 2.5.1, Aer 0.17.2"
    )
    set_run_font(run, size=10, color=MUTED)
    metadata.add_run("\nPrepared 24 July 2026")


def add_code_block(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.35)
    paragraph.paragraph_format.right_indent = Inches(0.35)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F7F9FB")
    properties.append(shading)
    run = paragraph.add_run(text)
    set_run_font(run, name="Consolas", size=9.3, color=NAVY)


def add_picture(
    document: Document,
    filename: str,
    *,
    width: float,
    alt_text: str,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    inline_shape = run.add_picture(str(TASK_DIR / filename), width=Inches(width))
    inline_shape._inline.docPr.set("descr", alt_text)


def build_report() -> None:
    backend = json.loads(
        (TASK_DIR / "task4_backend_validation.json").read_text(encoding="utf-8")
    )
    optimization = json.loads(
        (TASK_DIR / "task4_quantum_optimization.json").read_text(encoding="utf-8")
    )
    classical = json.loads(
        (TASK_DIR / "task4_classical_baseline.json").read_text(encoding="utf-8")
    )
    benchmark = json.loads(
        (TASK_DIR / "task4_quantum_classical_benchmark.json").read_text(
            encoding="utf-8"
        )
    )
    maximum_backend_difference = max(
        record["absolute_aer_quokka_difference"]
        for record in backend["records"]
    )
    q_metrics = benchmark["quantum"]["metrics"]
    c_metrics = benchmark["classical"]["metrics"]

    document = Document()
    configure_document(document)
    add_cover(document)

    add_linked_section(document)
    document.add_heading("Executive summary", level=1)
    add_body(
        document,
        "The course-supplied quantum machine-learning starter was converted into a "
        "local, deterministic workflow that retains its four-qubit basis encoder, "
        "recursive RY-RY-CX classifier, one-bit output, mean absolute-error objective, "
        "and SPSA training approach. The final notebook identifies exactly where data "
        "enters the circuit, explains the quantum model independently of its optimizer, "
        "records every optimization iteration and cumulative time, and validates one "
        "fixed parameter vector on exact statevector, Aer, and real Quokka backends."
    )
    add_callout(
        document,
        "Key result",
        f"Both implementations classify all four supplied images correctly at the "
        f"shared 0.5 threshold. The quantum model's full-dataset probability MAE is "
        f"{q_metrics['full']['mean_absolute_error']:.4f}; the classical model's is "
        f"{c_metrics['full']['mean_absolute_error']:.4f}. The classical path is faster "
        "and uses no quantum circuit or shots. Because the dataset has four samples and "
        "one test item, this is not evidence of generalisation or quantum advantage.",
    )
    document.add_heading("Assessment coverage", level=2)
    add_table(
        document,
        ["Requirement", "Implemented evidence", "Primary location"],
        [
            [
                "Analyse circuit/input/problem",
                "Pixel-qubit map, basis encoding, pair tree, gates, output and diagnostics",
                "Sections 1–2 / notebook",
            ],
            [
                "Iteration metric and time plots",
                "61 trace points: initial point plus 60 SPSA updates",
                "Section 3 / CSV / PNGs",
            ],
            [
                "No-quantum local version",
                "Two classical orientation features plus logistic regression",
                "Section 4 / classical_ml.py",
            ],
            [
                "Compare efficiency/effectiveness",
                "Shared split, threshold, accuracy, MAE, timings, circuits and shots",
                "Section 5 / benchmark JSON",
            ],
            [
                "Notebook and Word report",
                "Executed notebook and rendered report with retained evidence",
                "Task 4 submission folder",
            ],
        ],
        [2520, 4410, 2430],
    )
    document.add_heading("Deliverable map", level=2)
    add_bullet(document, "Executed analysis: Task_4_Quantum_ML.ipynb.")
    add_bullet(document, "Comprehensive report: Task_4_Quantum_ML_Report.docx.")
    add_bullet(
        document,
        "Machine-readable evidence: backend validation, optimization trace, classical "
        "baseline, and quantum/classical benchmark.",
    )
    add_bullet(
        document,
        "Review visuals: input dataset, circuit, two optimization plots, and the "
        "four-panel comparison.",
    )

    add_linked_section(document)
    add_section_heading(document, "1. Problem formulation and data")
    add_body(
        document,
        "The task is binary classification of non-uniform 2 x 2 binary images generated "
        "inside the starter notebook [1]. Vertical stripe patterns are label 0 and "
        "horizontal bar patterns are label 1. Removing the all-zero and all-one cases "
        "leaves four images. Each row-major flattened pixel vector x = (x0,x1,x2,x3) "
        "contains exactly two ones."
    )
    add_caption(document, "Figure 1. Complete generated Task 4 dataset and labels.")
    add_picture(
        document,
        "task4_dataset.png",
        width=5.7,
        alt_text=(
            "Four 2 by 2 binary images: two vertical stripes labelled zero and two "
            "horizontal bars labelled one."
        ),
    )
    document.add_heading("1.1 Deterministic split", level=2)
    add_body(
        document,
        "The starter's 75/25 random split is retained but isolated in a NumPy Generator "
        "with seed 802. Training indices are [3, 0, 2], with labels [1, 0, 1]; test "
        "index [1] has label 0. The training set therefore contains both classes. A "
        "single-item test partition is useful for reproducibility but statistically weak."
    )
    add_table(
        document,
        ["Sample", "Flattened pixels", "Class", "Split"],
        [
            ["stripe_1", "[1, 0, 1, 0]", "0", "Training"],
            ["stripe_2", "[0, 1, 0, 1]", "0", "Test"],
            ["bar_1", "[1, 1, 0, 0]", "1", "Training"],
            ["bar_2", "[0, 0, 1, 1]", "1", "Training"],
        ],
        [1900, 3300, 1700, 2460],
    )
    document.add_heading("1.2 Reproducibility boundary", level=2)
    add_body(
        document,
        "Local source modules replace Colab file upload, a hard-coded HTTP endpoint, "
        "and mutable global randomness. Dependencies are locked, backend configuration "
        "is read from an ignored environment file, and saved Quokka responses contain "
        "no credential. Remote evidence is retained rather than regenerated during "
        "ordinary notebook execution."
    )

    add_linked_section(document)
    add_section_heading(document, "2. Quantum circuit and input analysis")
    document.add_heading("2.1 Exact data-input location", level=2)
    add_body(
        document,
        "The data enters only through the basis-encoding loop. Pixel xi = 1 appends X "
        "to qubit qi; a zero leaves that qubit in |0>. The six angles are trainable "
        "model parameters, not sample features. Qubit and classical-bit labels follow "
        "Qiskit's documented circuit conventions [2]."
    )
    add_code_block(
        document,
        "for qubit, pixel in enumerate(sample):\n"
        "    if pixel > 0.5:\n"
        "        circuit.x(qubit)",
    )
    document.add_heading("2.2 Recursive parameterized model", level=2)
    add_body(
        document,
        "A post-order recursion yields pairs (q0,q1), (q2,q3), then (q1,q3). Each "
        "block applies RY(theta_k) to the first qubit, RY(theta_k+1) to the second, and "
        "CX from first to second. RY is a single-qubit Y-axis rotation [3]; CX makes "
        "the second branch conditional on the first. Three blocks require six angles. "
        "The final prediction is P(c0=1), obtained by measuring q3 into c0."
    )
    add_caption(
        document,
        "Figure 2. Example classifier for stripe_1 at zero trainable angles.",
    )
    add_picture(
        document,
        "task4_example_circuit.png",
        width=6.2,
        alt_text=(
            "Four-qubit classifier circuit for stripe one: X basis encoding, three "
            "recursive RY RY CX blocks, and measurement of qubit three."
        ),
    )
    document.add_heading("2.3 Zero-angle diagnostic", level=2)
    add_body(
        document,
        "At zero angles the RY gates are identities and the three CX operations form a "
        "parity tree. Output q3 becomes x0 XOR x1 XOR x2 XOR x3. Every supplied image "
        "has even parity, so exact statevector evaluation returns P(1)=0 for all four "
        "samples. This demonstrates that the untrained circuit cannot separate the "
        "classes; learned rotations are essential."
    )
    document.add_heading("2.4 Backend portability check", level=2)
    add_body(
        document,
        "One circuit constructor supports exact statevector probability, local "
        "shot-based Aer, and Quokka submission through OpenQASM 2. Qiskit documents "
        "qasm2.dumps as its circuit-to-string interoperability route [5]. AerSimulator "
        "is a configurable circuit simulator [4]. The real Quokka check uses a fixed "
        "angle vector and 256 shots per sample; it is validation, not training."
    )
    add_table(
        document,
        ["Sample", "Exact P(1)", "Aer P(1)", "Quokka P(1)", "|Aer–Quokka|"],
        [
            [
                record["sample_name"],
                f"{record['exact_probability_one']:.4f}",
                f"{record['aer']['probability_one']:.4f}",
                f"{record['quokka']['probability_one']:.4f}",
                f"{record['absolute_aer_quokka_difference']:.4f}",
            ]
            for record in backend["records"]
        ],
        [1900, 1750, 1750, 1960, 2000],
    )
    add_callout(
        document,
        "Backend finding",
        f"The maximum absolute Aer–Quokka difference is "
        f"{maximum_backend_difference:.6f}, equal to four counts out of 256. This is "
        "small relative to finite-shot uncertainty and does not indicate a systematic "
        "backend disagreement for the checked circuit.",
    )

    add_linked_section(document)
    add_section_heading(document, "3. Optimization evidence")
    document.add_heading("3.1 Objective and protocol", level=2)
    add_body(
        document,
        "The starter objective is training mean absolute error: average over samples of "
        "|P_theta(c0=1 | xi) - yi|. Seeded SPSA performs 60 updates on local Aer with "
        "256 shots per circuit. Each update evaluates a positive perturbation, a "
        "negative perturbation, and the updated parameter vector. SPSA estimates a "
        "multivariate search direction from simultaneous perturbations [7]. This report "
        "uses it as the supplied classical optimizer and does not treat it as part of "
        "the quantum circuit analysis."
    )
    add_table(
        document,
        ["Measure", "Recorded value", "Interpretation"],
        [
            [
                "Training MAE",
                f"{optimization['initial_objective_mae']:.4f} → "
                f"{optimization['final_objective_mae']:.4f}",
                "Lower is better",
            ],
            [
                "Best point",
                f"{optimization['best_objective_mae']:.4f} at iteration "
                f"{optimization['best_iteration']}",
                "All points retained",
            ],
            [
                "Objective calls",
                str(optimization["objective_evaluations"]),
                "Initial plus three per update",
            ],
            [
                "Circuit executions",
                str(optimization["circuit_executions"]),
                "Three training samples per objective",
            ],
            [
                "Training shots",
                f"{optimization['total_shots']:,}",
                "256 shots per circuit",
            ],
            [
                "Wall time",
                f"{optimization['elapsed_seconds']:.3f} s",
                "Machine-dependent local timing",
            ],
        ],
        [2350, 2680, 4330],
    )
    document.add_heading("3.2 Per-iteration metric", level=2)
    add_caption(
        document,
        "Figure 3. Training mean absolute error at iteration 0 and 60 SPSA updates.",
    )
    add_picture(
        document,
        "task4_quantum_metric_by_iteration.png",
        width=5.55,
        alt_text=(
            "Line chart of quantum training mean absolute error decreasing from "
            "0.5638 at iteration zero to 0.2904 at iteration sixty."
        ),
    )
    add_body(
        document,
        "The saved metric decreases from 0.5638 to 0.2904. Small flat or rising steps "
        "would be expected for a sampled stochastic search; retaining the complete CSV "
        "prevents selective reporting. In this seeded run the final point is also the "
        "best recorded point."
    )
    document.add_heading("3.3 Per-iteration cumulative time", level=2)
    add_caption(
        document,
        "Figure 4. Cumulative local Aer training time for the same trace.",
    )
    add_picture(
        document,
        "task4_quantum_time_by_iteration.png",
        width=5.55,
        alt_text=(
            "Line chart of cumulative local Aer training time increasing approximately "
            "linearly across sixty SPSA iterations."
        ),
    )
    add_body(
        document,
        "Cumulative time is monotonic and approximately linear because every update has "
        "the same objective-call structure. The absolute slope is specific to this "
        "machine; the retained circuit and shot counts are the more portable resource "
        "measure."
    )

    add_section_heading(document, "4. Fully classical no-circuit version")
    document.add_heading("4.1 Modification", level=2)
    add_body(
        document,
        "The local baseline is deliberately isolated in mse802/classical_ml.py. It "
        "imports NumPy and scikit-learn only and does not import Qiskit, create QASM, "
        "construct a quantum circuit, run a simulator, call Quokka, or consume shots. "
        "It converts each image to the mean change between rows and the mean change "
        "between columns, then fits balanced logistic regression on the same training "
        "indices. Scikit-learn documents LogisticRegression as its regularized "
        "classifier implementation [6]."
    )
    add_table(
        document,
        ["Sample", "Vertical change", "Horizontal change", "Label", "P(class 1)"],
        [
            [
                name,
                f"{features[0]:.1f}",
                f"{features[1]:.1f}",
                str(label),
                f"{probability:.4f}",
            ]
            for name, features, label, probability in zip(
                ("stripe_1", "stripe_2", "bar_1", "bar_2"),
                classical["all_orientation_features"],
                classical["full_dataset"]["labels"],
                classical["full_dataset"]["probability_one"],
                strict=True,
            )
        ],
        [1900, 2000, 2360, 1300, 1800],
    )
    add_callout(
        document,
        "No-quantum check",
        "The evidence records quantum_circuits = 0, quantum_shots = 0, and "
        "quantum_backend = null. An automated source check also rejects Qiskit, "
        "QuantumCircuit, QuokkaClient, or qasm2 references in the classical module.",
    )
    document.add_heading("4.2 Why the baseline works", level=2)
    add_body(
        document,
        "For a vertical stripe, rows are identical but adjacent columns differ, giving "
        "[vertical change, horizontal change] = [0,1]. A horizontal bar gives [1,0]. "
        "These features expose the data-generation rule directly, so all four points are "
        "linearly separable. This makes the baseline strong but also unusually tailored "
        "to the tiny synthetic dataset."
    )
    document.add_heading("4.3 Effectiveness", level=2)
    add_body(
        document,
        f"The classical model obtains training, test, and full-dataset accuracy of "
        f"{classical['training']['accuracy']:.1f}, "
        f"{classical['test']['accuracy']:.1f}, and "
        f"{classical['full_dataset']['accuracy']:.1f}. Its full-dataset MAE is "
        f"{classical['full_dataset']['mean_absolute_error']:.4f}. These values show "
        "correct operation on the supplied points, not performance on a broad image "
        "distribution."
    )

    add_section_heading(document, "5. Quantum/classical comparison")
    document.add_heading("5.1 Fair comparison protocol", level=2)
    add_body(
        document,
        "Both models use seed 802, the same three training indices and one test index, "
        "the same labels, and the same 0.5 decision threshold. Accuracy captures the "
        "thresholded decision; MAE retains probability confidence. Efficiency includes "
        "training wall time, median warmed four-sample inference time, and quantum "
        "resource counts. Quokka network/queue time is excluded because optimization "
        "was intentionally local."
    )
    add_caption(
        document,
        "Figure 5. Shared effectiveness and local efficiency benchmark.",
    )
    add_picture(
        document,
        "task4_quantum_classical_comparison.png",
        width=5.3,
        alt_text=(
            "Four-panel benchmark: both models have accuracy one; the classical model "
            "has lower mean absolute error and lower training and inference time."
        ),
    )
    add_table(
        document,
        ["Measure", "Quantum", "Classical", "Finding"],
        [
            [
                "Training accuracy",
                f"{q_metrics['training']['accuracy']:.3f}",
                f"{c_metrics['training']['accuracy']:.3f}",
                "Tie",
            ],
            [
                "Test accuracy",
                f"{q_metrics['test']['accuracy']:.3f}",
                f"{c_metrics['test']['accuracy']:.3f}",
                "One correct test item each",
            ],
            [
                "Full accuracy",
                f"{q_metrics['full']['accuracy']:.3f}",
                f"{c_metrics['full']['accuracy']:.3f}",
                "Tie",
            ],
            [
                "Full MAE",
                f"{q_metrics['full']['mean_absolute_error']:.4f}",
                f"{c_metrics['full']['mean_absolute_error']:.4f}",
                "Classical lower",
            ],
            [
                "Training time",
                f"{benchmark['quantum']['efficiency']['training_seconds']:.4f} s",
                f"{benchmark['classical']['efficiency']['training_seconds']:.4f} s",
                "Classical faster locally",
            ],
            [
                "Circuits / shots",
                f"{optimization['circuit_executions']} / "
                f"{optimization['total_shots']:,}",
                "0 / 0",
                "Classical uses no quantum resource",
            ],
        ],
        [2150, 2050, 2050, 3110],
    )
    document.add_heading("5.2 Interpretation", level=2)
    add_body(
        document,
        "At threshold 0.5 both pipelines make the same four correct decisions. The "
        "classical probabilities are substantially closer to their labels and its "
        "runtime/resource requirements are lower. For this narrowly defined problem, "
        "the classical approach is therefore more effective by MAE and more efficient. "
        "The quantum implementation remains useful as an auditable demonstration of "
        "basis encoding, parameterized RY-CX blocks, finite-shot training, and portable "
        "execution across local and remote backends."
    )
    add_callout(
        document,
        "Claim boundary",
        "No quantum advantage is claimed. A four-sample benchmark with one test item "
        "cannot establish generalisation, scalability, or hardware performance.",
    )

    add_section_heading(document, "6. Verification, limitations, and conclusion")
    document.add_heading("6.1 Verification", level=2)
    add_table(
        document,
        ["Check", "Result"],
        [
            ["Dataset and split", "Exact four images; deterministic [3,0,2] / [1] split"],
            ["Circuit schedule", "Pairs (0,1), (2,3), (1,3); six parameters; q3 readout"],
            ["Zero-angle diagnostic", "Exact output probability 0 for all four images"],
            ["Backend check", "Four real Quokka records; maximum difference retained"],
            ["Optimization", "61 ordered trace rows; nondecreasing cumulative time"],
            ["No-circuit baseline", "No quantum imports/circuits/shots; all points classified"],
            ["Repository tests", "29 tests passed at report generation"],
            ["Notebook", "Executed top-to-bottom from a fresh kernel"],
        ],
        [3000, 6360],
    )
    document.add_heading("6.2 Limitations", level=2)
    add_bullet(
        document,
        "The dataset has four synthetic samples and the test result is one decision.",
    )
    add_bullet(
        document,
        "The classical features directly encode the known generation rule and may not "
        "transfer to richer images.",
    )
    add_bullet(
        document,
        "Aer is ideal unless a noise model is added; the Quokka run checks only one "
        "fixed vector at 256 shots per sample.",
    )
    add_bullet(
        document,
        "SPSA results depend on hyperparameters, initialization, finite shots, and seed; "
        "only one fully retained run is reported.",
    )
    add_bullet(
        document,
        "Wall-clock timings are machine-specific and exclude remote network/queue delay.",
    )
    document.add_heading("6.3 Conclusion", level=2)
    add_body(
        document,
        "The completed Task 4 workflow meets the assessment requirements with an "
        "explicitly analysed quantum circuit, reproducible backend validation, complete "
        "iteration/time evidence, a genuinely no-circuit local version, and a common "
        "efficiency/effectiveness benchmark. The retained notebook, source modules, raw "
        "records, CSV traces, figures, and automated tests make every reported value "
        "reviewable. Within the severe data limitations, the classical solution is the "
        "practical winner; the quantum solution demonstrates the intended circuit and "
        "execution concepts without overstating its result."
    )

    document.add_heading("References", level=1)
    references = [
        "[1] MSE802 Quantum Computing, “Quantum_ML_AS2.ipynb,” course-supplied "
        "Assessment 2 starter notebook, 2026.",
        "[2] IBM Quantum, “Quantum circuit model,” Qiskit SDK documentation. "
        "https://quantum.cloud.ibm.com/docs/api/qiskit/circuit "
        "(accessed 24 July 2026).",
        "[3] IBM Quantum, “RYGate,” Qiskit SDK API documentation. "
        "https://quantum.cloud.ibm.com/docs/api/qiskit/1.3/"
        "qiskit.circuit.library.RYGate (accessed 24 July 2026).",
        "[4] Qiskit Development Team, “AerSimulator,” Qiskit Aer documentation. "
        "https://qiskit.github.io/qiskit-aer/stubs/qiskit_aer.AerSimulator.html "
        "(accessed 24 July 2026).",
        "[5] IBM Quantum, “OpenQASM 2 and the Qiskit SDK.” "
        "https://quantum.cloud.ibm.com/docs/guides/interoperate-qiskit-qasm2 "
        "(accessed 24 July 2026).",
        "[6] scikit-learn developers, “LogisticRegression,” scikit-learn API. "
        "https://scikit-learn.org/stable/modules/generated/"
        "sklearn.linear_model.LogisticRegression.html (accessed 24 July 2026).",
        "[7] J. C. Spall, “Multivariate stochastic approximation using a simultaneous "
        "perturbation gradient approximation,” IEEE Transactions on Automatic Control, "
        "vol. 37, no. 3, pp. 332–341, 1992.",
        "[8] M. A. Nielsen and I. L. Chuang, Quantum Computation and Quantum Information, "
        "10th anniversary ed. Cambridge: Cambridge University Press, 2010.",
    ]
    for reference in references:
        paragraph = document.add_paragraph(reference)
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(6)

    add_linked_section(document)
    document.add_heading("Appendix A — Evidence inventory", level=1)
    add_table(
        document,
        ["Artefact", "Purpose"],
        [
            ["Task_4_Quantum_ML.ipynb", "Executed end-to-end analysis"],
            ["task4_backend_validation.json", "Real fixed-vector Quokka comparison"],
            ["task4_quantum_optimization*.json/.csv/.png", "Training parameters, trace and plots"],
            ["task4_classical_baseline.json", "No-circuit model evidence"],
            ["task4_quantum_classical_benchmark.*", "Common protocol, results and figure"],
            ["task4_example_circuit.qasm/.png", "Exact circuit representation"],
            ["SOURCE_NOTE.md / README.md", "Provenance and reproduction instructions"],
            ["Task_4_Quantum_ML_Report.docx", "Comprehensive assessment report"],
        ],
        [4380, 4980],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
