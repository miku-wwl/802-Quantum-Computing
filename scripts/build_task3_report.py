"""Generate the Task 3 Word report with deterministic document styling."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TASK_DIR = ROOT / "submission" / "Task_3_Quantum_Tic_Tac_Toe"
OUTPUT = TASK_DIR / "Task_3_Quantum_Tic_Tac_Toe_Report.docx"

# Documents skill preset: standard_business_brief.
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
NAVY = RGBColor(0x20, 0x37, 0x50)
MUTED = RGBColor(0x5A, 0x65, 0x70)
LIGHT_BLUE = "EAF2F8"
LIGHT_GREY = "F2F4F7"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_section_geometry(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)


def add_linked_section(document: Document):
    """Add a consistently formatted content section linked to prior furniture."""

    section = document.add_section(WD_SECTION.NEW_PAGE)
    set_section_geometry(section)
    for header_footer_part in (
        section.header,
        section.even_page_header,
        section.first_page_header,
        section.footer,
        section.even_page_footer,
        section.first_page_footer,
    ):
        header_footer_part.is_linked_to_previous = True
    return section


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for margin_name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = margins.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError("table columns must total 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    width.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    width.set(qn("w:type"), "dxa")
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths_dxa:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(value))
        grid.append(grid_column)

    for row in table.rows:
        for cell, value in zip(row.cells, widths_dxa, strict=True):
            cell.width = Inches(value / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            cell_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            cell_width.set(qn("w:w"), str(value))
            cell_width.set(qn("w:type"), "dxa")


def repeat_header(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    row_properties.append(repeat)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def configure_document(document: Document) -> None:
    document.settings.odd_and_even_pages_header_footer = False
    section = document.sections[0]
    set_section_geometry(section)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style_name in ("List Bullet", "List Number"):
        style = styles[list_style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = MUTED
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(5)
    caption.paragraph_format.keep_with_next = True

    for header_part in (
        section.header,
        section.even_page_header,
        section.first_page_header,
    ):
        header = header_part.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_run = header.add_run("MSE802  •  ASSESSMENT 2  •  TASK 3")
        set_run_font(header_run, size=8.5, color=MUTED, bold=True)

    for footer_part in (
        section.footer,
        section.even_page_footer,
        section.first_page_footer,
    ):
        footer = footer_part.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("Quantum Tic-Tac-Toe  |  Page ")
        set_run_font(footer_run, size=8.5, color=MUTED)
        add_page_number(footer)

    document.core_properties.title = "Quantum Tic-Tac-Toe — Task 3 Report"
    document.core_properties.subject = "MSE802 Assessment 2"
    document.core_properties.author = "MSE802 student submission"
    document.core_properties.keywords = "quantum computing, Qiskit, Tic-Tac-Toe"


def add_body(document: Document, text: str, *, bold_lead: str | None = None):
    paragraph = document.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        paragraph.add_run(text[len(bold_lead) :])
    else:
        paragraph.add_run(text)
    return paragraph


def add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text, style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_section_heading(document: Document, text: str, *, level: int = 1):
    """Add a major heading; Word's keep-with-next style prevents orphans."""

    return document.add_heading(text, level=level)


def add_callout(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.10
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_BLUE)
    properties.append(shading)
    lead = paragraph.add_run(f"{label}  ")
    set_run_font(lead, color=DARK_BLUE, bold=True)
    paragraph.add_run(text)


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_dxa: list[int],
):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, LIGHT_GREY)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=9.5, color=NAVY, bold=True)
    repeat_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_run_font(run, size=9.2)
    set_table_geometry(table, widths_dxa)
    return table


def add_cover(document: Document) -> None:
    for _ in range(5):
        document.add_paragraph()
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("TECHNICAL INVESTIGATION REPORT")
    set_run_font(run, size=10, color=BLUE, bold=True)
    kicker.paragraph_format.space_after = Pt(18)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Quantum Tic-Tac-Toe")
    set_run_font(run, size=30, color=NAVY, bold=True)
    title.paragraph_format.space_after = Pt(8)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Code completion, gameplay dynamics,\nquantum-circuit generation, and gate analysis"
    )
    set_run_font(run, size=15, color=DARK_BLUE)
    subtitle.paragraph_format.space_after = Pt(30)

    rule = document.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = rule.add_run("MSE802 Quantum Computing  •  Assessment 2  •  Task 3")
    set_run_font(run, size=11, color=MUTED, bold=True)
    rule.paragraph_format.space_after = Pt(98)

    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = metadata.add_run("Local implementation: Python 3.11, Qiskit 2.5.1, Aer 0.17.2")
    set_run_font(run, size=10, color=MUTED)
    metadata.add_run("\nPrepared 24 July 2026")


def build_report() -> None:
    evidence = json.loads((TASK_DIR / "task3_game_evidence.json").read_text(encoding="utf-8"))
    document = Document()
    configure_document(document)
    add_cover(document)

    add_linked_section(document)
    document.add_heading("Executive summary", level=1)
    add_body(
        document,
        "The supplied Quantum Tic-Tac-Toe notebook was completed and converted from a "
        "Google Colab-only prototype into a local, reproducible Jupyter application. "
        "Every move now produces an explicit Qiskit gate, all eight winning triples are "
        "implemented, measurement order is handled correctly, SWAP uses a reliable "
        "two-cell interaction, and Replay creates a clean nine-qubit circuit."
    )
    add_callout(
        document,
        "Key result",
        "Four saved games and seven Task 3 automated tests confirm the gate semantics, "
        "measurement mapping, score calculation, clean reset, and interface flow. "
        "The complete project test suite contains 14 passing tests.",
    )
    document.add_heading("Assessment coverage", level=2)
    add_table(
        document,
        ["Requirement", "Implemented evidence", "Location"],
        [
            [
                "Complete missing code",
                "Not, O, X, SWAP, eight wins, measurement, reset",
                "mse802/tictactoe.py",
            ],
            [
                "Play multiple games",
                "Four seeded games with move logs and resolved boards",
                "task3_game_evidence.json",
            ],
            [
                "Analyse code and gameplay",
                "Sections 2–6 plus executed notebook",
                "This report / Task 3 notebook",
            ],
            [
                "Explain circuit generation",
                "State preparation, move-to-gate mapping, copied measurement circuit",
                "Sections 3–4",
            ],
            [
                "Describe every gate",
                "Reset, H, RY, X, SWAP, barrier, measurement",
                "Section 4",
            ],
        ],
        [2600, 4300, 2460],
    )
    document.add_heading("Deliverable map", level=2)
    add_bullet(document, "Executed notebook: Task_3_Quantum_Tic_Tac_Toe.ipynb.")
    add_bullet(document, "Game implementation: mse802/tictactoe.py; interface: mse802/tictactoe_ui.py.")
    add_bullet(document, "Reproducible evidence: four QASM files, four circuit PNGs, summary PNG, and JSON.")
    add_bullet(document, "Automated checks: tests/test_tictactoe.py.")

    add_section_heading(document, "1. Purpose and completed architecture")
    add_body(
        document,
        "The educational purpose is to make circuit construction visible through a "
        "familiar game. The pre-measurement board is not classical ownership data: it is "
        "a nine-qubit state assembled by the player's sequence of reversible gates. "
        "Measurement ends a round, converts each qubit to O or X, and evaluates the "
        "classical winning lines."
    )
    document.add_heading("1.1 Adaptation of the starter", level=2)
    add_body(
        document,
        "The starter mixed package installation, game rules, simulator calls, and a "
        "Colab-specific grid in one notebook. The local version separates those concerns:"
    )
    add_bullet(document, "The notebook explains and demonstrates the assessment work.")
    add_bullet(document, "Board owns the quantum circuit, validated moves, measurement, and scoring.")
    add_bullet(document, "TicTacToeWidget owns only portable controls and presentation.")
    add_bullet(document, "Evidence and notebook-builder scripts make generated artefacts repeatable.")
    add_bullet(document, "Pytest exercises the model independently of manual interaction.")
    document.add_heading("1.2 Runtime flow", level=2)
    add_table(
        document,
        ["Stage", "State change", "Observable result"],
        [
            ["New round", "Reset then H on q0…q8", "Nine unresolved |+⟩ cells"],
            ["Choose move", "Append X, RY, or SWAP", "Move log and circuit update"],
            ["Measure", "Copy circuit; measure qi → ci", "One c8…c0 bitstring"],
            ["Post-process", "Reverse string; 0→O, 1→X", "Board cells 0…8"],
            ["Score", "Evaluate eight triples", "Separate X and O line counts"],
            ["Replay", "Discard old circuit; prepare new board", "No stale gate or score"],
        ],
        [1700, 3900, 3760],
    )
    add_callout(
        document,
        "Design decision",
        "Measurement is appended to a circuit copy. The unmeasured move circuit remains "
        "available for inspection, and repeated UI drawing cannot accidentally add "
        "measurement instructions.",
    )

    add_section_heading(document, "2. Quantum state and circuit generation")
    document.add_heading("2.1 Initial board", level=2)
    add_body(
        document,
        "A new round creates QuantumCircuit(9, 9). Each qubit is reset to |0⟩ and then "
        "receives H, producing |+⟩ = (|0⟩ + |1⟩)/√2. Before any move, the joint state is "
        "|+⟩⊗9 and every board pattern has probability 1/512 under ideal independent "
        "measurement. Reset is retained from the starter even though a new simulator "
        "circuit begins in |0⟩; it makes the preparation intention explicit."
    )
    document.add_heading("2.2 How gameplay generates a circuit", level=2)
    add_body(
        document,
        "Board.apply_move validates the selected cell and dispatches the operation to a "
        "Qiskit circuit method. Gates are appended in click order, so later actions act "
        "on the quantum state produced by earlier actions. The move history is metadata "
        "for review; only the circuit determines the measured outcome."
    )
    code = document.add_paragraph()
    code.paragraph_format.left_indent = Inches(0.35)
    code.paragraph_format.right_indent = Inches(0.35)
    code.paragraph_format.space_before = Pt(6)
    code.paragraph_format.space_after = Pt(10)
    properties = code._p.get_or_add_pPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), "F7F9FB")
    properties.append(shade)
    run = code.add_run(
        'O   → circuit.ry(-π/2, cell)\n'
        'X   → circuit.ry(+π/2, cell)\n'
        'Not → circuit.x(cell)\n'
        'SWAP→ circuit.swap(first_cell, second_cell)'
    )
    set_run_font(run, name="Consolas", size=9.5, color=NAVY)
    document.add_heading("2.3 Measurement and Qiskit ordering", level=2)
    add_body(
        document,
        "The measurable copy maps qi to ci for i = 0…8 and executes exactly one shot on "
        "AerSimulator. Qiskit prints bit n−1 at the left of a string and bit 0 at the "
        "right [2]. Therefore c8…c0 is reversed before assigning the first character to "
        "board cell 0. The implementation maps 0 to O and 1 to X, matching the starter "
        "comments that O rotates toward |0⟩ and X toward |1⟩."
    )
    add_callout(
        document,
        "Failure prevented",
        "Without the reversal, moves on cell 0 would appear on cell 8. The automated "
        "bit-order test fixes both endpoint cells and verifies the seven middle cells.",
    )

    add_section_heading(document, "3. Quantum gates and instructions")
    add_body(
        document,
        "This section covers every instruction used to generate or resolve the game "
        "circuit. Quantum gates are unitary and reversible; reset and measurement are "
        "non-unitary operations."
    )
    add_table(
        document,
        ["Instruction", "Action in this game", "Key characteristic"],
        [
            ["reset", "Prepare every square as |0⟩", "Non-unitary initialization"],
            ["H", "Create unresolved |+⟩ square", "Equal 0/1 amplitudes"],
            ["RY(−π/2)", "Set an untouched |+⟩ square to O", "Y-axis rotation to |0⟩"],
            ["RY(+π/2)", "Set an untouched |+⟩ square to X", "Y-axis rotation to |1⟩"],
            ["Pauli-X", "Not: exchange 0 and 1 amplitudes", "X² = I; bit flip"],
            ["SWAP", "Exchange two complete cell states", "|a,b⟩ → |b,a⟩ [3]"],
            ["barrier", "Visually separate preparation/resolution", "No state change"],
            ["measure", "Resolve qi into classical ci", "Born-rule probabilistic collapse"],
        ],
        [1700, 3900, 3760],
    )
    document.add_heading("3.1 Hadamard and open squares", level=2)
    add_body(
        document,
        "H = (1/√2)[[1, 1], [1, −1]]. Acting on |0⟩ gives |+⟩. An untouched cell is "
        "therefore genuinely unresolved, not an empty classical value. A one-shot "
        "measurement returns O or X with equal ideal probability."
    )
    document.add_heading("3.2 RY rotations for O and X", level=2)
    add_body(
        document,
        "RY(θ) = [[cos(θ/2), −sin(θ/2)], [sin(θ/2), cos(θ/2)]]. Substitution gives "
        "RY(−π/2)|+⟩ = |0⟩ and RY(+π/2)|+⟩ = |1⟩. These identities make O and X "
        "deterministic on a fresh square. Repeated rotations still compose as quantum "
        "operations; the interface intentionally does not pretend a text label is the "
        "current measured state."
    )
    document.add_heading("3.3 Pauli-X as Not", level=2)
    add_body(
        document,
        "X = [[0, 1], [1, 0]] exchanges |0⟩ and |1⟩. Thus Not converts a prepared O "
        "to X and X to O. On an untouched |+⟩ cell, X|+⟩ = |+⟩, so Not has no "
        "observable effect—exactly the gameplay description in the starter."
    )

    add_section_heading(document, "3.4 SWAP, barrier, and measurement", level=2)
    add_body(
        document,
        "SWAP is a two-qubit Clifford gate with mapping |a,b⟩ → |b,a⟩ [3]. It exchanges "
        "the full amplitudes of two cells, not their interface labels. The local UI "
        "collects two different cells before appending one SWAP; selecting the same cell "
        "cancels the operation."
    )
    add_body(
        document,
        "Barriers group the initial preparation and final resolution visually. They do "
        "not change amplitudes and may constrain compiler reordering. Measurement is "
        "different: it returns a classical bit according to squared amplitudes and ends "
        "the quantum portion of the round. AerSimulator is a configurable noisy-circuit "
        "simulator; this assessment uses its default ideal CPU simulation [4]."
    )
    add_caption(document, "Figure 1. Game 3 measured circuit: preparation, SWAP/Not moves, and resolution.")
    document.add_picture(str(TASK_DIR / "game_3_swap.png"), width=Inches(6.45))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading("3.5 Gate-order implications", level=2)
    add_body(
        document,
        "Gate order matters. In Game 3, cell 0 is prepared O and cell 1 X; SWAP makes "
        "cell 0 X and cell 1 O; Not then turns cell 1 into X. Preparing cell 2 as X "
        "completes the top row. Moving Not before SWAP would flip a different logical "
        "state and could change the winner."
    )

    add_section_heading(document, "4. Gameplay dynamics and scoring")
    document.add_heading("4.1 Interaction", level=2)
    add_body(
        document,
        "The player chooses a move and clicks a board cell. For SWAP, the first click is "
        "highlighted and the second click completes the gate. Measure resolves all cells "
        "in one shot. Replay resets the circuit, history, score, pending SWAP selection, "
        "button labels, and button colours."
    )
    document.add_heading("4.2 Winning conditions", level=2)
    add_body(
        document,
        "The eight triples are rows (0,1,2), (3,4,5), (6,7,8); columns (0,3,6), "
        "(1,4,7), (2,5,8); and diagonals (0,4,8), (2,4,6). The scorer counts all "
        "completed lines independently for both symbols. Unlike ordinary alternating "
        "Tic-Tac-Toe, a quantum resolution can produce simultaneous or multiple wins."
    )
    add_table(
        document,
        ["Game", "Resolved board (rows separated by /)", "X lines", "O lines"],
        [
            [
                record["id"].replace("_", " ").title(),
                "/".join(
                    "".join(record["result"]["board"][start : start + 3])
                    for start in (0, 3, 6)
                ),
                str(record["result"]["wins_x"]),
                str(record["result"]["wins_o"]),
            ]
            for record in evidence["games"]
        ],
        [1900, 4760, 1350, 1350],
    )
    document.add_heading("4.3 Why labels are provisional", level=2)
    add_body(
        document,
        "Before measurement, a button label records operations such as O, N, or S; it "
        "does not calculate a classical owner. This is essential for superposition and "
        "gate composition. After measurement, buttons display only resolved O/X values "
        "and the score reports completed lines."
    )
    add_callout(
        document,
        "Interpretation",
        "The game is an educational circuit builder, not a claim of quantum advantage. "
        "Its value is the direct connection between a user action, a gate, and a "
        "measurable state transition.",
    )

    add_section_heading(document, "5. Multiple-game evidence")
    add_caption(document, "Figure 2. Four submitted one-shot game outcomes.")
    document.add_picture(
        str(TASK_DIR / "task3_four_game_summary.png"),
        width=Inches(5.4),
    )
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading("5.1 Games 1–3: analytical checks", level=2)
    add_body(
        document,
        "Game 1 resolves as OOO/XXX/XXX, yielding one O row and two X rows. Game 2 "
        "starts O everywhere and applies Not on 0, 4, and 8, yielding XOO/OXO/OOX and "
        "one X diagonal. Game 3 resolves XXX/OOO/OOO after SWAP and Not, yielding one X "
        "row and two O rows. All three results exactly match the gate algebra."
    )
    document.add_heading("5.2 Game 4: probabilistic resolution", level=2)
    game_four = evidence["games"][3]["result"]
    add_body(
        document,
        f"Game 4 appends no player gate. Seed 805 produced "
        f"{''.join(game_four['board'])}, with {game_four['wins_x']} X lines and "
        f"{game_four['wins_o']} O lines. The seed makes the saved run reproducible, "
        "but it does not make this board a deterministic physical prediction. A new "
        "seed can produce any of 512 ideal board strings."
    )

    document.add_heading("6. Verification and engineering quality", level=1)
    document.add_heading("6.1 Automated checks", level=2)
    add_table(
        document,
        ["Check", "Assertion"],
        [
            ["Gate dispatch", "O/X/Not/SWAP append ry/ry/x/swap in order"],
            ["Deterministic semantics", "O→0, X→1, and O+Not→1"],
            ["State exchange", "SWAP exchanges prepared O and X cells"],
            ["Winning triples", "Exactly eight unique triples; all-X board scores eight"],
            ["Bit order", "c8…c0 reverses to cells 0…8"],
            ["Measurement isolation", "Stored game circuit receives no measure gate"],
            ["UI flow", "Two-click SWAP and Replay reset behave correctly"],
        ],
        [2700, 6660],
    )
    add_body(
        document,
        "At evidence generation, all seven Task 3 tests passed and the complete repository "
        "suite reported 14 passed. The notebook was also executed from top to bottom "
        "without a cell error. Seeds 802–805 are stored with the game records."
    )
    document.add_heading("6.2 Reproducibility", level=2)
    add_bullet(document, "Dependencies are locked by pyproject.toml and uv.lock.")
    add_bullet(document, "The notebook is generated by scripts/build_task3_notebook.py.")
    add_bullet(document, "Game artefacts are regenerated by scripts/generate_task3_evidence.py.")
    add_bullet(document, "QASM retains the exact measured circuit for each saved game.")
    add_bullet(document, "JSON retains versions, backend, seeds, moves, depth, operations, and outcomes.")
    document.add_heading("6.3 Limitations", level=2)
    add_body(
        document,
        "Aer is ideal unless a noise model is supplied, so the evidence does not measure "
        "hardware error. One shot is correct for resolving one game but is insufficient "
        "for estimating probabilities. Most played circuits contain no entangling gate: "
        "SWAP exchanges states but does not itself create entanglement from product "
        "inputs. Finally, repeated O/X rotations are allowed as quantum operations even "
        "when their classical-looking labels may suggest simple ownership."
    )

    add_section_heading(document, "7. Conclusion")
    add_body(
        document,
        "The completed application meets the Task 3 requirements as a local, interactive, "
        "and auditable quantum program. The circuit is generated directly from gameplay; "
        "its preparation, moves, measurement, endianness conversion, and scoring are "
        "explicit. Multiple saved games demonstrate deterministic rotation, Not, SWAP, "
        "simultaneous wins, and an unresolved superposition board. Automated tests and "
        "retained QASM/JSON evidence support the analysis rather than relying on screenshots "
        "alone."
    )
    document.add_heading("References", level=1)
    references = [
        "[1] MSE802 Quantum Computing, “Quantum_Tic_Tac_Toe__AS2.ipynb,” "
        "course-supplied Assessment 2 starter notebook, 2026.",
        "[2] IBM Quantum, “Bit-ordering in the Qiskit SDK,” IBM Quantum Documentation. "
        "https://quantum.cloud.ibm.com/docs/en/guides/bit-ordering (accessed 24 July 2026).",
        "[3] IBM Quantum, “SwapGate,” Qiskit SDK API documentation. "
        "https://quantum.cloud.ibm.com/docs/api/qiskit/qiskit.circuit.library.SwapGate "
        "(accessed 24 July 2026).",
        "[4] Qiskit Development Team, “AerSimulator,” Qiskit Aer 0.17 documentation. "
        "https://qiskit.github.io/qiskit-aer/stubs/qiskit_aer.AerSimulator.html "
        "(accessed 24 July 2026).",
        "[5] M. A. Nielsen and I. L. Chuang, Quantum Computation and Quantum Information, "
        "10th anniversary ed. Cambridge: Cambridge University Press, 2010.",
    ]
    for reference in references:
        paragraph = document.add_paragraph(reference)
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(6)
    document.add_heading("Appendix A — Evidence inventory", level=1)
    add_table(
        document,
        ["Artefact", "Purpose"],
        [
            ["Task_3_Quantum_Tic_Tac_Toe.ipynb", "Executed explanation and interactive game"],
            ["game_1_rows … game_4_open (.qasm/.png)", "Exact circuits and diagrams"],
            ["task3_game_evidence.json", "Machine-readable moves, metrics, and results"],
            ["task3_four_game_summary.png", "Visual comparison of resolved boards"],
            ["SOURCE_NOTE.md", "Starter provenance and adaptation boundary"],
            ["Task_3_Quantum_Tic_Tac_Toe_Report.docx", "Comprehensive assessment report"],
        ],
        [3900, 5460],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
